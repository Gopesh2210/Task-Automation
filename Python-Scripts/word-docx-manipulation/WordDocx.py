'''
    This script is used for word-docx file manipulation
'''	
import docx

'''
    READ DOCX
'''
# loading docx document
d = docx.Document('demo.docx')

# saving second paragraph text
p2 = d.paragraphs[1].text
print("Second Paragraph : ",p2)

# acquuring the different run texts
# runs : different text styles

p = d.paragraphs[1]

r1 = p.runs[1].text
r2 = p.runs[2].text
r3 = p.runs[3].text

print("RUN 1 : " + str(r1))
print("RUN 2 : " + str(r2))
print("RUN 3 : " + str(r3))


# changing the style
title = d.paragraphs[3]
run = p.runs[3]
run.underline = True
run.bold = True
run.text = 'underlined and bolded'
title.style = 'Title'

d.save('edited-demo.docx')

'''
    WRITE DOCX
'''
# creating a new docx
doc = docx.Document()

# adding a paragraph
doc.add_paragraph('New Paragraph Added.')
para = doc.paragraphs[0]

# adding a new run
para.add_run('This is a new run')
new_run = para.runs[1]

# changing run style
new_run.bold = True

doc.save('new-demo.docx')